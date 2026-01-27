import os
from pathlib import Path
from typing import Union, Tuple
import logging

logger = logging.getLogger(__name__)


class DocumentLoaderError(Exception):
    """Custom exception for document loading errors."""
    pass


class DocumentLoader:
    """Loads and extracts raw text from PDF and DOC/DOCX files."""
    
    SUPPORTED_FORMATS = {'.pdf', '.doc', '.docx'}
    
    @staticmethod
    def load(file_path: Union[str, Path]) -> str:
        """
        Load and extract raw text from a document.
        
        Args:
            file_path: Path to PDF, DOC, or DOCX file
            
        Returns:
            Raw text extracted from the document
            
        Raises:
            DocumentLoaderError: If file format is not supported or loading fails
        """
        file_path = Path(file_path)
        
        # Validate file exists
        if not file_path.exists():
            raise DocumentLoaderError(f"File not found: {file_path}")
        
        # Validate file format
        file_ext = file_path.suffix.lower()
        if file_ext not in DocumentLoader.SUPPORTED_FORMATS:
            raise DocumentLoaderError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(DocumentLoader.SUPPORTED_FORMATS)}"
            )
        
        try:
            if file_ext == '.pdf':
                return DocumentLoader._load_pdf(file_path)
            elif file_ext == '.docx':
                return DocumentLoader._load_docx(file_path)
            elif file_ext == '.doc':
                return DocumentLoader._load_doc(file_path)
        except DocumentLoaderError:
            raise
        except Exception as e:
            raise DocumentLoaderError(f"Error loading {file_ext} file: {str(e)}")
    
    @staticmethod
    def _load_pdf(file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise DocumentLoaderError(
                "PyPDF2 not installed. Install with: pip install PyPDF2"
            )
        
        text = []
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                num_pages = len(pdf_reader.pages)
                
                if num_pages == 0:
                    logger.warning(f"PDF has no pages: {file_path}")
                    return ""
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                
                logger.info(f"Extracted text from {num_pages} pages in {file_path.name}")
        
        except Exception as e:
            raise DocumentLoaderError(f"Failed to read PDF: {str(e)}")
        
        return "\n\n".join(text)
    
    @staticmethod
    def _load_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise DocumentLoaderError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        
        text = []
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            logger.info(f"Extracted text from DOCX: {file_path.name}")
        
        except Exception as e:
            raise DocumentLoaderError(f"Failed to read DOCX: {str(e)}")
        
        return "\n".join(text)
    
    @staticmethod
    def _load_doc(file_path: Path) -> str:
        """Extract text from older DOC format."""
        try:
            from docx import Document
        except ImportError:
            raise DocumentLoaderError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        
        try:
            # Attempt to open as DOCX (some .doc files are actually DOCX format)
            doc = Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            logger.info(f"Extracted text from DOC: {file_path.name}")
            return "\n".join(text)
        
        except Exception as e:
            # If standard approach fails, provide helpful error
            raise DocumentLoaderError(
                f"Failed to read DOC file. Note: Older binary .doc files may require "
                f"additional libraries. Error: {str(e)}"
            )
    
    @staticmethod
    def load_batch(file_paths: list, skip_errors: bool = False) -> Tuple[dict, list]:
        """
        Load multiple documents.
        
        Args:
            file_paths: List of file paths
            skip_errors: If True, skip files that fail to load and continue
            
        Returns:
            Tuple of (success_dict, error_list)
            - success_dict: {file_path: extracted_text}
            - error_list: List of (file_path, error_message) tuples
        """
        results = {}
        errors = []
        
        for file_path in file_paths:
            try:
                text = DocumentLoader.load(file_path)
                results[str(file_path)] = text
            except DocumentLoaderError as e:
                error_msg = str(e)
                logger.error(f"Failed to load {file_path}: {error_msg}")
                errors.append((str(file_path), error_msg))
                
                if not skip_errors:
                    raise
        
        return results, errors
